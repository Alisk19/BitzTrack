import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    # Load the original letterhead template
    doc = docx.Document('dist/assets/3d1 invoice.docx')

    # Add some spacing
    doc.add_paragraph()
    
    # Add Invoice Metadata
    p = doc.add_paragraph()
    p.add_run("Invoice ID: ").bold = True
    p.add_run("{invoiceId}\n")
    p.add_run("Date: ").bold = True
    p.add_run("{date}\n")
    p.add_run("Customer: ").bold = True
    p.add_run("{customerName}")
    
    doc.add_paragraph()

    # Add Table
    table = doc.add_table(rows=2, cols=6)
    pass
    
    # Headers
    hr = table.rows[0].cells
    hr[0].text = 'Sr.No'
    hr[1].text = 'Description'
    hr[2].text = 'Material'
    hr[3].text = 'Quantity'
    hr[4].text = 'Rate'
    hr[5].text = 'Total'
    
    # Make headers bold
    for cell in hr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data row with docxtemplater tags
    dr = table.rows[1].cells
    dr[0].text = '{#items}{srNo}'
    dr[1].text = '{description}'
    dr[2].text = '{material}'
    dr[3].text = '{quantity}'
    dr[4].text = '{rate}'
    dr[5].text = '{total}{/items}'
    
    doc.add_paragraph()
    
    # Add Totals
    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_total.add_run("Grand Total: ").bold = True
    p_total.add_run("Rs. {totalAmount}")
    
    doc.add_paragraph()
    
    p_words = doc.add_paragraph()
    p_words.add_run("Amount in Words: ").bold = True
    p_words.add_run("{amountInWords}")
    
    # Save to public folder so the app can fetch it properly
    doc.save('public/3d1 invoice.docx')
    print("Template generation complete!")

try:
    create_template()
except Exception as e:
    print("Error:", e)
